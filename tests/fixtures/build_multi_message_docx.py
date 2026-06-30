"""Build tests/fixtures/csg_multi_message.docx — numbered titles + KV meta tables."""

from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


def build_multi_message_docx(path: Path) -> Path:
    if Document is None:
        raise ImportError("python-docx required")

    doc = Document()
    doc.add_heading("基础协议报文定义", level=1)

    doc.add_paragraph("1. 查询通信延时时长")
    t0 = doc.add_table(rows=2, cols=2)
    t0.rows[0].cells[0].text = "功能码"
    t0.rows[0].cells[1].text = "03"
    t0.rows[1].cells[0].text = "数据标识"
    t0.rows[1].cells[1].text = "E8030304"
    t1 = doc.add_table(rows=1, cols=3)
    t1.rows[0].cells[0].text = "字段"
    t1.rows[0].cells[1].text = "长度"
    t1.rows[0].cells[2].text = "说明"
    r = t1.add_row().cells
    r[0].text = "dest_addr"
    r[1].text = "6字节"
    r[2].text = "通信目的地址"

    doc.add_paragraph("2. 查询从节点信息")
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "功能码"
    t2.rows[0].cells[1].text = "03"
    t2.rows[1].cells[0].text = "数据标识"
    t2.rows[1].cells[1].text = "E8030306"
    t2b = doc.add_table(rows=1, cols=3)
    t2b.rows[0].cells[0].text = "字段"
    t2b.rows[0].cells[1].text = "长度"
    t2b.rows[0].cells[2].text = "说明"
    r2 = t2b.add_row().cells
    r2[0].text = "start_slave_index"
    r2[1].text = "2字节"
    r2[2].text = "从节点起始序号"

    doc.add_paragraph("3. 其他报文（无数据标识）")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


if __name__ == "__main__":
    out = Path(__file__).parent / "csg_multi_message.docx"
    build_multi_message_docx(out)
    print(f"written {out}")
