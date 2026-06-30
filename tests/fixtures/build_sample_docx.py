"""Build tests/fixtures/csg_sample.docx for doc_parser integration tests."""

from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


def build_sample_docx(path: Path) -> Path:
    if Document is None:
        raise ImportError("python-docx required")

    doc = Document()
    doc.add_heading("AFN03 DI=E80304F5 查询设备类型", level=2)
    doc.add_paragraph("下行请求，无地址域。")

    t1 = doc.add_table(rows=1, cols=3)
    t1.rows[0].cells[0].text = "字段"
    t1.rows[0].cells[1].text = "长度"
    t1.rows[0].cells[2].text = "说明"
    r1 = t1.add_row().cells
    r1[0].text = "设备类型"
    r1[1].text = "2字节"
    r1[2].text = "00H：单相表；01H：三相表；02H：采集器"

    doc.add_heading("AFN03 DI=E80304F6 查询通信延时", level=2)
    doc.add_paragraph("下行请求。")

    t2 = doc.add_table(rows=1, cols=3)
    t2.rows[0].cells[0].text = "字段"
    t2.rows[0].cells[1].text = "长度"
    t2.rows[0].cells[2].text = "说明"
    r2 = t2.add_row().cells
    r2[0].text = "timeout"
    r2[1].text = "2字节"
    r2[2].text = "超时时间(秒)"

    r3 = t2.add_row().cells
    r3[0].text = "保留"
    r3[1].text = "4字节"
    r3[2].text = "厂家私有透明数据"

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


if __name__ == "__main__":
    out = Path(__file__).parent / "csg_sample.docx"
    build_sample_docx(out)
    print(f"written {out}")
